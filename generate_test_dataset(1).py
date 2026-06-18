import json
from pathlib import Path

from docx import Document


ROOT = Path("/Users/ngc4594/Documents/智能系统skill")

SOURCES = {
    "cycling": Path("/Users/ngc4594/Downloads/访谈记录.docx"),
    "concert_eire": Path("/Users/ngc4594/Downloads/20260330160005-manggo-3月30日16_00-16_30-逐字稿文本-1.docx"),
    "concert_pumpkin": Path("/Users/ngc4594/Downloads/20260330150106-转写_李筱晴的快速会议-逐字稿文本-1.docx"),
}


def read_docx(path: Path) -> str:
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            vals = [c.text.strip() for c in row.cells if c.text.strip()]
            if vals:
                parts.append(" | ".join(vals))
    return "\n".join(parts)


def anonymize(text: str) -> str:
    replacements = {
        "徐春琳": "访谈者A",
        "李筱晴": "访谈者B",
        "芝士": "访谈者C",
        "天泽泽泽泽BIGL9": "P-CYCLE-01",
        "小琳": "P-CYCLE-01",
        "饶同学": "P-CYCLE-02",
        "佘同学": "P-CYCLE-03",
        "小陈": "P-CYCLE-04",
        "小宇": "P-CYCLE-05",
        "小阳": "P-CYCLE-06",
        "Eire": "P-CONCERT-01",
        "PumpkinTJQ": "P-CONCERT-02",
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)
    return text


def main() -> None:
    cycling_text = anonymize(read_docx(SOURCES["cycling"]))
    concert_eire_text = anonymize(read_docx(SOURCES["concert_eire"]))
    concert_pumpkin_text = anonymize(read_docx(SOURCES["concert_pumpkin"]))

    test_data = {
        "dataset_name": "user_research_synthesis_skill_test_data",
        "version": "1.0",
        "created_for": "user-research-synthesis skill final evaluation",
        "language": "zh-CN",
        "privacy_note": "Participant names and interviewer names have been anonymized. Raw wording is otherwise preserved for qualitative coding tests.",
        "test_cases": [
            {
                "id": "T001_CONCERT_EXPERIENCE",
                "title": "演唱会体验访谈",
                "method": "semi_structured_interview_transcripts",
                "difficulty": "medium",
                "data_source": "real_course_research_data_anonymized",
                "research_question": "用户在演唱会或演出体验中如何记录、回看、分享和收集内容？有哪些可转化为平台功能的机会点？",
                "analysis_boundary": "聚焦演唱会/演出前后信息查找、现场记录、朋友圈/群聊分享、回看行为、物料收集、打卡活动参与意愿；不分析票务购买流程和艺人偏好本身。",
                "source_files": [
                    str(SOURCES["concert_eire"]),
                    str(SOURCES["concert_pumpkin"]),
                ],
                "participants": [
                    {
                        "id": "P-CONCERT-01",
                        "profile": "演唱会参与者，近期去过大型演唱会，偏好拍摄歌手舞台和应援物。",
                        "source_file": SOURCES["concert_eire"].name,
                    },
                    {
                        "id": "P-CONCERT-02",
                        "profile": "演出/演唱会参与者，更偏生活记录、朋友圈和朋友群分享。",
                        "source_file": SOURCES["concert_pumpkin"].name,
                    },
                ],
                "input_text": "=== 受访者 P-CONCERT-01 ===\n"
                + concert_eire_text
                + "\n\n=== 受访者 P-CONCERT-02 ===\n"
                + concert_pumpkin_text,
                "expected_output_types": [
                    "segments",
                    "code_groups",
                    "themes",
                    "contradictions",
                    "insights",
                    "markdown_report",
                ],
            },
            {
                "id": "T002_CYCLING_EXPERIENCE",
                "title": "骑行体验访谈",
                "method": "mixed_interview_transcript_and_summaries",
                "difficulty": "hard",
                "data_source": "real_course_research_data_anonymized",
                "research_question": "用户为什么低频骑行或放弃骑行？骑行 App 或智能硬件如何降低门槛、提升乐趣并支持绿色出行？",
                "analysis_boundary": "聚焦骑行障碍、骑行前准备、骑行中体验、骑行后复盘、游戏化激励、社群陪伴、环保反馈和二手装备信任；不分析专业竞技训练和自行车硬件制造。",
                "source_files": [str(SOURCES["cycling"])],
                "participants": [
                    {
                        "id": "P-CYCLE-01",
                        "profile": "22岁女性，低频骑行，曾有中学骑行经历，现多用电动车。",
                    },
                    {
                        "id": "P-CYCLE-02",
                        "profile": "23岁女性，骑行小白，一个月骑行2-3次。",
                    },
                    {
                        "id": "P-CYCLE-03",
                        "profile": "24岁男性工程师，放弃骑行，当前更偏电动车。",
                    },
                    {
                        "id": "P-CYCLE-04",
                        "profile": "22岁男性上班族，目前很少骑行，重视效率和成本。",
                    },
                    {
                        "id": "P-CYCLE-05",
                        "profile": "26岁男性，骑行资深用户，通勤每天骑行。",
                    },
                ],
                "excluded_content": [
                    {
                        "label": "P-CYCLE-06",
                        "reason": "文档中仅有标题和版本号，没有可分析正文。",
                    }
                ],
                "input_text": "=== 骑行访谈数据 ===\n" + cycling_text,
                "expected_output_types": [
                    "segments",
                    "code_groups",
                    "themes",
                    "contradictions",
                    "insights",
                    "markdown_report",
                ],
            },
        ],
    }

    reference_answers = {
        "dataset_name": "user_research_synthesis_skill_reference_answers",
        "version": "1.0",
        "scoring_note": "These are human reference answers for evaluating theme coverage, insight quality, contradiction handling, and actionability. Exact wording does not need to match, but outputs should cover the same evidence and design implications.",
        "references": [
            {
                "test_case_id": "T001_CONCERT_EXPERIENCE",
                "human_analysis_summary": "人工分析认为，演唱会用户的记录行为围绕现场情绪、阶段反应、朋友圈生活记录和实体/电子物料收集展开。Skill 应识别记录、回看、分享、攻略、打卡和物料收集等主题，但行动建议需要人工进一步聚焦。",
                "expected_themes": [
                    {
                        "id": "C-T01",
                        "name": "现场记录",
                        "description": "用户会在入场、开场、爱豆出场、互动、临场发挥等情绪节点记录照片、视频或文字。",
                        "must_include_evidence": [
                            "一开始要记录一下环境",
                            "爱豆出场的时候",
                            "和嘉宾互动、听众互动",
                            "live 图可以进入当场的瞬间",
                        ],
                    },
                    {
                        "id": "C-T02",
                        "name": "回看复盘",
                        "description": "用户会在发布前检查，也会在事后通过朋友圈、相册或回忆功能回看内容。",
                        "must_include_evidence": [
                            "发之前我会检查一下",
                            "会回看一下自己以前玩过什么去过哪里",
                            "上新的时候回看一下",
                        ],
                    },
                    {
                        "id": "C-T03",
                        "name": "私域分享",
                        "description": "朋友圈和朋友群是主要分享场域，记录生活比获取点赞更重要。",
                        "must_include_evidence": [
                            "主要是记录生活",
                            "没有人点赞评论依然会发",
                            "把自己拍到的全发群里面",
                        ],
                    },
                    {
                        "id": "C-T04",
                        "name": "模板与打卡",
                        "description": "用户对一键模板、人生四格、顺路打卡和可打印纪念物有兴趣，但要求操作简单、路线不绕。",
                        "must_include_evidence": [
                            "操作简单",
                            "一键生成",
                            "不要让我再去绕路",
                            "可以免费打印出来",
                        ],
                    },
                    {
                        "id": "C-T05",
                        "name": "攻略信息",
                        "description": "用户会在小红书、大众点评等平台查攻略、座位视野、场馆信息、周边吃喝和售卖集合点。",
                        "must_include_evidence": [
                            "小红书会比较多一点",
                            "查攻略",
                            "周边吃喝玩乐",
                            "场馆布置",
                        ],
                    },
                    {
                        "id": "C-T06",
                        "name": "物料收集",
                        "description": "用户有票根、海报、玩偶、电子票根等收集需求，但收集后不一定频繁回看。",
                        "must_include_evidence": [
                            "我一定会收集",
                            "收集完之后我就放在那儿，也没怎么看",
                            "票根、海报、玩偶",
                            "电子票根",
                        ],
                    },
                ],
                "expected_core_insights": [
                    {
                        "statement": "记录发生在关键情绪和阶段节点，而不只是随机拍照。",
                        "confidence": "high",
                    },
                    {
                        "statement": "回看行为常由发布检查、朋友圈复盘和阶段性情绪反应触发。",
                        "confidence": "high",
                    },
                    {
                        "statement": "用户偏好真实、低成本的记录方式，live 图和原相机比复杂模板更稳妥。",
                        "confidence": "medium",
                    },
                    {
                        "statement": "物料与票根有纪念价值，但收集后低频回顾，适合设计轻量召回机制。",
                        "confidence": "medium",
                    },
                    {
                        "statement": "打卡活动必须顺路、低难度、可获得明确纪念回报，才会提升参与意愿。",
                        "confidence": "medium",
                    },
                ],
                "expected_action_opportunities": [
                    "演唱会前提供座位视野、入场路线、周边吃喝和集合点聚合页。",
                    "演唱会中/后提供一键生成真实感 live 记录卡或朋友圈图文模板。",
                    "设计顺路打卡任务，避免用户为打卡额外绕路。",
                    "将电子票根、照片、文字故事、定位和物料整理为可回看的演出记忆页。",
                    "为物料收集或点评打卡提供小奖励，但不应干扰用户沉浸式观看。",
                ],
                "known_skill_comparison": {
                    "theme_coding": "主题编码结果应基本与问题设置一致，主要需要修改编码名称。",
                    "insight_coverage": "应覆盖记录在出场前、回看因为阶段反应、记录形式以 live 为主、有收集物料癖好等核心洞察；若补充“收集后很少回顾”“愿意参加点评打卡活动”，需由人工判断是否过度延展。",
                    "actionability": "行动建议可能偏弱，需要人工进一步聚焦活动机制或产品方案。",
                },
            },
            {
                "test_case_id": "T002_CYCLING_EXPERIENCE",
                "human_analysis_summary": "人工分析将骑行体验按 Journey Map 组织为骑行前、骑行中、骑行后和长期互动，核心设计愿景是把骑行重新定义为低门槛的短途城市探索/都市娱乐方式，而不是高强度运动。",
                "expected_themes": [
                    {
                        "id": "B-T01",
                        "name": "骑行前门槛",
                        "description": "路线规划、目的地选择、装备准备和结伴安排会提高小白启动成本。",
                        "must_include_evidence": [
                            "路况复杂",
                            "路线规划",
                            "装备",
                            "组队",
                        ],
                    },
                    {
                        "id": "B-T02",
                        "name": "骑行中安全与疲劳",
                        "description": "用户在骑行中担心路况、安全、出汗、疲劳和操作不便。",
                        "must_include_evidence": [
                            "极端天气",
                            "路况复杂",
                            "骑行中看手机危险",
                            "疲惫",
                        ],
                    },
                    {
                        "id": "B-T03",
                        "name": "数据与成长",
                        "description": "里程、速度、轨迹、健康和减碳数据是骑行后复盘与成就感的重要来源。",
                        "must_include_evidence": [
                            "码表",
                            "轨迹",
                            "里程",
                            "健康",
                            "减碳",
                        ],
                    },
                    {
                        "id": "B-T04",
                        "name": "游戏化与探索乐趣",
                        "description": "虚拟小人、任务、积分和路线探索有吸引力，但交互必须低干扰。",
                        "must_include_evidence": [
                            "虚拟 IP",
                            "积分奖励",
                            "语音鼓励",
                            "不能一直看手机",
                        ],
                    },
                    {
                        "id": "B-T05",
                        "name": "社群陪伴",
                        "description": "用户需要陪伴、指导、分层社群和约伴活动，但微信仍承担熟人沟通。",
                        "must_include_evidence": [
                            "组队",
                            "陪伴",
                            "指导",
                            "阶段分层",
                            "微信",
                        ],
                    },
                    {
                        "id": "B-T06",
                        "name": "环保转化",
                        "description": "环保不是所有人的强动机，更适合作为骑后可视化成就、积分或福利。",
                        "must_include_evidence": [
                            "环保态度悲观",
                            "减碳数据",
                            "绿色出行活动",
                            "积分兑换",
                        ],
                    },
                    {
                        "id": "B-T07",
                        "name": "二手与装备信任",
                        "description": "二手装备交易存在卖家信誉、品质、售后和卫生风险，需要骑友经验或平台保障。",
                        "must_include_evidence": [
                            "骗子",
                            "假货",
                            "描述不符",
                            "售后",
                            "骑友建议",
                        ],
                    },
                ],
                "expected_core_insights": [
                    {
                        "statement": "降低启动门槛，帮助骑行小白迈出第一步。",
                        "confidence": "high",
                    },
                    {
                        "statement": "让骑行过程变得更有趣，对抗身体疲劳。",
                        "confidence": "high",
                    },
                    {
                        "statement": "通过数据可视化了解自身成长。",
                        "confidence": "high",
                    },
                    {
                        "statement": "用持续奖励机制激励绿色出行。",
                        "confidence": "medium",
                    },
                    {
                        "statement": "骑行被替代不是兴趣消失，而是当前方式更省心。",
                        "confidence": "medium",
                    },
                    {
                        "statement": "安全路线预判是低频用户重新骑行的前置条件。",
                        "confidence": "high",
                    },
                    {
                        "statement": "游戏化应服务骑后反馈，而不是干扰骑行过程。",
                        "confidence": "high",
                    },
                ],
                "expected_action_opportunities": [
                    "新手友好的路线与目的地规划。",
                    "新手友好 Check list。",
                    "自然交互方式，如语音、震动、低干扰提醒。",
                    "创新玩法机制，把城市转化为可探索、可收集的体验。",
                    "App 数据报告，展示里程、轨迹、健康和环保成果。",
                    "社交平台分享功能。",
                    "App 社区模块，按入门、中度、资深分层。",
                    "环保积分兑换系统。",
                    "二手装备验货清单、骑友推荐和平台保障。",
                ],
                "known_skill_comparison": {
                    "theme_coding": "Skill 能识别骑行替代、安全路线、数据反馈、游戏化激励、社群陪伴、环保转化、二手交易信任等 7 个主题。",
                    "insight_coverage": "Skill 更偏横向主题归纳，人工分析更偏 Journey Map 和设计机会转化。",
                    "actionability": "Skill 建议可作为功能清单，但最终愿景仍需人工综合为“低门槛短途城市探索/都市娱乐方式”。",
                },
            },
        ],
        "evaluation_rubric": {
            "theme_coverage": "Output should cover at least 70% of expected themes for each test case.",
            "evidence_traceability": "Each top insight should include 2-3 supporting quotes or clear source references.",
            "contradiction_handling": "Output should identify conflicting attitudes instead of forcing a single conclusion.",
            "actionability": "Recommendations should specify target user, context, action, and expected effect.",
            "human_review_need": "Outputs may be accepted as first-round research synthesis, but final design vision requires human synthesis.",
        },
    }

    (ROOT / "test_data.json").write_text(
        json.dumps(test_data, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    (ROOT / "reference_answers.json").write_text(
        json.dumps(reference_answers, ensure_ascii=False, indent=2), encoding="utf-8"
    )


if __name__ == "__main__":
    main()
